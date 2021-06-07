# pip is the package installer for python3. The command "py -m pip install"
# installs the packages where in "-m pip" executes pip using the latest python interpreter installed in windows.

# Package to get access the working directory
import os
# Package to split and modify the search input
import re

### packages for reading excel files
from openpyxl import Workbook
from openpyxl import load_workbook

### packages for writing word documents
from docx import Document
from docx.shared import Cm
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX

### packages for webscraping
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
import time
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.action_chains import ActionChains

#Setting up selenium driver
#Install the corresponding version of chromedriver on https://sites.google.com/a/chromium.org/chromedriver/downloads and place the file in a map. Store the path to the file in the PATH variable below.
'''
option = webdriver.ChromeOptions()
option.add_argument('headless')
PATH = "C:\Program Files (x86)\chromedriver.exe"
driver = webdriver.Chrome(PATH, options=option)
'''
# package to read the pdf files
from tabula import read_pdf
from tabulate import tabulate
import tabula

# Package to open and read the text files
import csv

# Package to access the API's
import requests as rq

# Package to access the clinvar API database
from eutils import Client

# Package to create the GUI
from tkinter import *

# Creating a function that allows for the creation of working hyperlinks

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

# Creating the graphical user interface to allow the user to interact with the script

# Creating a click function that saves the user input
def click():
    global brca_input
    global sequence_type
    global search
    brca_input = gene_entry.get()
    sequence_type = sequence_entry.get()
    search = search_entry.get()
    print(("Searching for {} {}").format(brca_input, search))



# Creating window to allow the user to interact with the script

window = Tk()
window.title("BRCA1/2 automatic consultation")

# Canvas to paint background color

canvas = Canvas(window, width=600, height=300, bg="SteelBlue3")
canvas.grid(columnspan=3, rowspan=4)

# Creating a title

Intro = Label(canvas, text="Automatic consultation of diverse databases required for BRCA1/2 gene identification", bg="lightblue")
Intro.grid(column=0,row=0)


# Creating the input labels for gene, sequence and search
gene = Label(canvas, text="Enter the gene: BRCA1/BRCA2:", bg="lightblue")
gene.grid(column=0, row=2)
gene_entry = Entry(canvas, width=20, bg="lightblue")
gene_entry.grid(row=2, column=1)

sequence = Label(canvas, text="Protein/nucleotide?:", bg="lightblue")
sequence.grid(column=0, row=3)
sequence_entry = Entry(canvas, width=20, bg="lightblue")
sequence_entry.grid(row=3, column=1)

search = Label(canvas, text="Enter sequence:", bg="lightblue")
search.grid(column=0, row=4)
search_entry = Entry(canvas, width=20, bg="lightblue")
search_entry.grid(row=4, column=1)

# Creating a submit button to confirm the search input
Button (canvas, text="SUBMIT", width=6, command=click) .grid(row=4, column=2)

brca_input = gene_entry.get()
sequence_type = sequence_entry.get()
search = search_entry.get()

# The window is a loop that stays on the screen
window.mainloop()

# Transforming the input
'''
brca_input = brca_input.upper()
sequence_type = sequence_type.upper()
search = search.upper()
'''
# To reduce the number of redundant searches there will be a range_check added. If you enter a nucleotide search that's just a number. Example: 5074
# The script will return as output only the results ranging from nucleotide variant location -4 to +4.
# Generating a list of numbers and converting them from integers into strings. The script will use this list as a comparison
# if the search output gives a nucleotide squence of +4 or -4 it will not be given as output because the list of disallowed numbers ranges from 4 to 9999 in this case
# this can be changed below.
pattern = list(range(4, 9999))
pattern_string = [str(number) for number in pattern]

# If the input is a number nucleotide sequence, the output should be compared to a range of values
# if the input is a specific SNP it shouldn't be compared to the list

if sequence_type == "PROTEIN":
    search_split = re.search("[0-9]+", search)
    search_mod = search_split.group()

# No range check will be done for example for c.3262C>T but there will be a range check for example 5074
elif sequence_type == "NUCLEOTIDE":
    if bool(re.search("[a-zA-Z]", search)) == True:
        search_split = re.split("[+, ., \-, >, _, a-zA-Z]", search)
        search_mod = search_split[2]
        range_check = False

    else:
        search_mod = search
        range_check = True

# Setting up the path to the location of the excel files and pdf file
# The script has to be in the same folder as the excel files
folder_location = os.getcwd()
folder_location = "C:/Users/jensv/Desktop/BRCA_-_prostate_cancer_and_ovarian_cancer_-_PARPi"
functional_categorization = "/2020 BRCA1 functional categorization of BRCA1 VUS - CCR.xlsx"
variants_excel = "/allenigmavariants_BICsubmission_2013-07-01.xlsx"
bicbnl = "/bicbnl 27  fh 2-10-13_nov2013.xlsx"
enigma = "/BRCA_Multifac_published data_ENIGMAwebsite 2015-03-27.xlsx"
pdf_lindor = "/Lindor 2012.pdf"
cmg_brca1 = "/BRCA1 CMG 05 2021.txt"
cmg_brca2 = "/BRCA2 CMG 05 2021.txt"
brca1_exchange = "/brca1_exchange"
brca2_exchange = "/brca2_exchange"
# PATH is the location of the chromedriver
PATH = "C:\Program Files (x86)\chromedriver.exe"

# Create word document
document = Document()
sections = document.sections
for section in sections:
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

# UZA CMG: priority: 1
# The script will open the text file depending on if the gene input is BRCA1/2
print("Searching in the UZA CMG file")
# Adding the table to the word document
document.add_heading("UZA CMG / BRCA Exchange / LOVD")
table_word = document.add_table(rows=1, cols=5)

table_word.style = 'Colorful List Accent 1'
hdr_cells = table_word.rows[0].cells
hdr_cells[0].text = 'Source'
hdr_cells[1].text = 'Gene'
hdr_cells[2].text = 'cDNA'
hdr_cells[3].text = 'Protein'
hdr_cells[4].text = 'Mut effect/Classification'

if brca_input == "BRCA1":

    full_path = folder_location + cmg_brca1
    with open(full_path) as csv_file:
        # The script will open the text file using csv.reader but using the tab ("\t") delimiter so it will iterate over each line of text separated by tabs.
        csv_reader = csv.reader(csv_file, delimiter="\t")
        for row in csv_reader:
            count = 0
            hit = False
            for cell in row:
                count = count + 1
                # The coding sequences are on the 8th tab
                if count == 8 and search in cell:
                    row_cells = table_word.add_row().cells
                    row_cells[0].text = "UZA CMG"
                    row_cells[1].text = brca_input
                    row_cells[2].text = cell
                    hit = True
                    count = count + 1
                    
                    
                # if it doesn't get a hit for the coding sequences it will check the protein sequence on the 9th tab
                elif count == 9 and search in cell:
                    row_cells = table_word.add_row().cells
                    row_cells[0].text = "UZA CMG"
                    row_cells[1].text = brca_input
                    row_cells[3].text = cell
                    hit = True
                    count = count + 1
                # The mut effect is on the 11th tab. If there is a hit for the coding sequences or the protein it will give the mut effect as output.
                if count == 11 and hit == True:
                    row_cells[4].text = cell
                    hit = False
                    count = count + 1
                    
             
        
# The files for BRCA1 and BRCA2 have the same structure so the loop can be repeated.
elif brca_input == "BRCA2":

    full_path = folder_location + cmg_brca2
    with open(full_path) as csv_file:
        csv_reader = csv.reader(csv_file, delimiter="\t")
        for row in csv_reader:
            count = 0
            hit = False
            for cell in row:
                count = count + 1
                if count == 8 and search in cell:
                    row_cells = table_word.add_row().cells
                    row_cells[0].text = "UZA CMG"
                    row_cells[1].text = brca_input
                    row_cells[2].text = cell
                    hit = True
                    count = count + 1
                elif count == 9 and search in cell:
                    row_cells = table_word.add_row().cells
                    row_cells[0].text = "UZA CMG"
                    row_cells[1].text = brca_input
                    row_cells[3].text = cell
                    hit = True
                    count = count + 1
                if count == 11 and hit == True:
                    row_cells[4].text = cell
                    hit = False
                    count = count + 1

# Searching BRCA Exchange database"
# priority: 2
# Setting up the driver. Headless mode is added so the sites don't constantly pop up when the script is running.
# The elements can be found by rightclicking on the page when visiting the website and clicking inspect
# The script finds the elements mostly by xpath
print("Searching the BRCA Exchange database")
option = webdriver.ChromeOptions()
option.add_argument('headless')
driver = webdriver.Chrome(PATH, options = option)
'''
self.chrome = webdriver.Chrome(ChromeDriverManager().install(), options=options)
self.chrome.execute_cdp_cmd("Page.setBypassCSP", {"enabled": True})
'''

# Try is used so the script doesn't break if it can't find results on this specific website
try:

    driver.get("https://brcaexchange.org/")

    # Looking for the searchbar and enter the search
    searchbar = driver.find_element_by_xpath("//input[@type='text']")
    searchbar.send_keys(brca_input + " " + search)
    searchbar.send_keys(Keys.RETURN)
    # wait for the page to load
    time.sleep(4)

    # Increasing the amount of hits per page
    page = driver.current_url + "&pageLength=100"
    driver.get(page)
    time.sleep(4)

    # Finding the result table
    table = driver.find_element_by_xpath("//*[@id='data-table-container']/div/table/tbody")
    
    tab_count = 1
    # The script will iterate over each cell in the table goes to the next tab if possible.
    while tab_count < 6:
        for row in table.find_elements_by_xpath("//tr"):
            count = 0
            # hit = False or hit = True is used to append the classification if there the user input is found on the site
            hit = False
            for table_data in row.find_elements_by_xpath('.//td'):
                # The script will iterate over the results and keep count of the current position
                count = count + 1
                if count == 2 and search in table_data.text:
                    # If there is a range check the script will perform it to reduce the number of output
                    if range_check == True:
                        cDNA = table_data.text
                        cDNA_split = re.split("[+, ., \-, >, _, a-zA-Z]", cDNA)
                        cDNA_split_plusmin = cDNA_split[3]

                        # The script will compare the hits with the list of numbers created earlier
                        if not any(i == cDNA_split_plusmin for i in pattern_string):
                            row_cells = table_word.add_row().cells
                            row_cells[0].text = "BRCA Exchange"
                            row_cells[1].text = brca_input
                            row_cells[2].text = table_data.text
                            hit = True
                    # If no range check is needed it will just add the coding sequence to the output
                    elif range_check == False:
                        row_cells = table_word.add_row().cells
                        row_cells[0].text = "BRCA Exchange"
                        row_cells[1].text = brca_input
                        row_cells[2].text = table_data.text
                        hit = True
                # If there was a hit on the coding sequence it will add the protein to the output
                if count == 3 and hit == True:
                    row_cells[3].text = table_data.text
                    hit = True
                # If the search is a protein and on this position it will add it to the output
                if count == 3 and search in table_data.text:
                    row_cells = table_word.add_row().cells
                    row_cells[0].text = "BRCA Exchange"
                    row_cells[1].text = brca_input
                    row_cells[3].text = table_data.text
                    hit = True
                # On position 6 is the classification, if there was a hit on coding sequence level or protein level it will add the classification to the output
                if count == 6 and hit == True:
                    row_cells[4].text = table_data.text
                    hit = False
        
        # The script will look for the button that direct to the next tab, if the button is clickable it will click on the tab and repeat previous loop
        next_button = driver.find_element_by_xpath("//*[@id='variant-search-row']/div[2]/ul/li[8]/a")
        enabled = next_button.is_enabled()
        if enabled == True:
            next_button.click()
            time.sleep(10)
            tab_count = tab_count + 1
            enabled = next_button.is_enabled()
        # If the button is not enabled (There are no more tabs of results) this loop will break
        else:
            p = document.add_paragraph("Results: ")    
            site_location = str(driver.current_url)
            add_hyperlink(p, 'Link', driver.current_url)
            break



#create and save a hyperlink to word
        p = document.add_paragraph("Results: ")    
        site_location = str(driver.current_url)
        add_hyperlink(p, 'Link', driver.current_url)
        driver.quit()

except Exception:
        pass


# Searching in the LOVD database: priority 3
print("Searching the LOVD database")

option = webdriver.ChromeOptions()
option.add_argument('headless')
driver = webdriver.Chrome(PATH)
#driver.execute_cdp_cmd("Page.setBypassCSP", {"enabled": True})


try:

    if brca_input == "BRCA1" and sequence_type == "NUCLEOTIDE":

        # The script will search for the searchbar on the site and enter the nucleotide search in it.

        driver.get("https://databases.lovd.nl/shared/variants/BRCA1/unique")
        searchbar = driver.find_element_by_name("search_VariantOnTranscript/DNA")
        searchbar.send_keys(search)
        searchbar.send_keys(Keys.RETURN)
    
        #time.sleep() is used to pause the script to give the website the time to load

        time.sleep(4)

        hit = False

        count = 0
        # Finding the table with results
        table = driver.find_element_by_xpath("//table[@id='viewlistTable_CustomVL_VOTunique_VOG_BRCA1']")
        for row in table.find_elements_by_xpath("//tr[@class='data']"):
            #The script will iterate over the results table in the page row by row
            count = 0
            for table_data in row.find_elements_by_xpath(".//td"):
                count = count + 1
                # If the search is in the cell on the 4th position it will add it to the output depending if a range check is needed or not
                if count == 4 and search in table_data.text:
                    if range_check == True:
                        cDNA = table_data.text
                        cDNA_split = re.split("[+, ., \-, >, _, a-zA-Z]", cDNA)
                        cDNA_split_plusmin = cDNA_split[3]

                        if not any(i == cDNA_split_plusmin for i in pattern_string):
                            row_cells = table_word.add_row().cells
                            row_cells[0].text = "LOVD"
                            row_cells[1].text = brca_input
                            row_cells[2].text = table_data.text
                            hit = True

                    elif range_check == False:
                        row_cells = table_word.add_row().cells
                        row_cells[0].text = "LOVD"
                        row_cells[1].text = brca_input
                        row_cells[2].text = table_data.text
                        hit = True
                # It will add the associated protein
                if count == 6 and hit == True:
                    row_cells[3].text = table_data.text
                    '''
                # If the user input was a protein it will add it to the output

                if count == 6 and search in table_data.text:
                    row_cells = table_word.add_row().cells
                    row_cells[0].text = "LOVD"
                    row_cells[1].text = brca_input
                    row_cells[3].text = table_data.text
                    hit = True
                    '''
                # If there was a hit it will add the classification
                if count == 9 and hit == True:
                    row_cells[4].text = table_data.text
                    hit = False               



        # It will append the results to the word document alongside with a hyperlink of the current url location
        p = document.add_paragraph("Results: ")    
        site_location = str(driver.current_url)
        add_hyperlink(p, 'Link', driver.current_url)

        driver.quit()

    elif brca_input == "BRCA2" and sequence_type == "NUCLEOTIDE":
        #If BRCA2 is given as input it will start the search on the BRCA2 tab of the website
        driver.get("https://databases.lovd.nl/shared/variants/BRCA2/unique")
        searchbar = driver.find_element_by_name("search_VariantOnTranscript/DNA")
        searchbar.send_keys(search)
        searchbar.send_keys(Keys.RETURN)
    
        time.sleep(4)

        hit = False
        count = 0
        table = driver.find_element_by_xpath("//table[@id='viewlistTable_CustomVL_VOTunique_VOG_BRCA2']")
        for row in table.find_elements_by_xpath("//tr[@class='data']"):
            count = 0
            for table_data in row.find_elements_by_xpath(".//td"):
                count = count + 1
                if count == 4 and search in table_data.text:
                    if range_check == True:
                        cDNA = table_data.text
                        cDNA_split = re.split("[+, ., \-, >, _, a-zA-Z]", cDNA)
                        cDNA_split_plusmin = cDNA_split[3]

                        if not any(i == cDNA_split_plusmin for i in pattern_string):
                            row_cells = table_word.add_row().cells
                            row_cells[0].text = "LOVD"
                            row_cells[1].text = brca_input
                            row_cells[2].text = table_data.text
                            hit = True

                    elif range_check == False:
                        row_cells = table_word.add_row().cells
                        row_cells[0].text = "LOVD"
                        row_cells[1].text = brca_input
                        row_cells[2].text = table_data.text
                        hit = True
                        print(table_data.text)

                if count == 6 and hit == True:
                    row_cells[3].text = table_data.text
                    print(table_data.text)

                '''
                elif count == 6 and search in table_data.text:
                    row_cells = table_word.add_row().cells
                    row_cells[0].text = "LOVD"
                    row_cells[1].text = brca_input
                    row_cells[3].text = table_data.text
                    hit = True
                '''
                if count == 9 and hit == True:
                    row_cells[4].text = table_data.text
                    hit = False
                    print(table_data.text)

        p = document.add_paragraph("Results: ")    
        site_location = str(driver.current_url)
        add_hyperlink(p, 'Link', driver.current_url)

        driver.quit()

    elif brca_input == "BRCA1" and sequence_type == "PROTEIN":
        driver.get("https://databases.lovd.nl/shared/variants/BRCA1/unique")
        searchbar = driver.find_element_by_name("search_VariantOnTranscript/Protein")
        searchbar.send_keys(search)
        searchbar.send_keys(Keys.RETURN)

        time.sleep(4)

        hit = False

        count = 0
        table = driver.find_element_by_xpath("//table[@id='viewlistTable_CustomVL_VOTunique_VOG_BRCA1']")
        for row in table.find_elements_by_xpath("//tr[@class='data']"):
            count = 0
            for table_data in row.find_elements_by_xpath(".//td"):
                count = count + 1
                if count == 6 and search in table_data.text:
                    row_cells = table_word.add_row().cells
                    row_cells[0].text = "LOVD"
                    row_cells[1].text = brca_input
                    row_cells[3].text = table_data.text
                    hit = True

                if count == 9 and hit == True:
                    row_cells[4].text = table_data.text
                    hit = False               



        # It will append the results to the word document alongside with a hyperlink of the current url location
        p = document.add_paragraph("Results: ")    
        site_location = str(driver.current_url)
        add_hyperlink(p, 'Link', driver.current_url)

        driver.quit()

    elif brca_input == "BRCA2" and sequence_type == "PROTEIN":
        #If BRCA2 is given as input it will start the search on the BRCA2 tab of the website
        driver.get("https://databases.lovd.nl/shared/variants/BRCA2/unique")
        searchbar = driver.find_element_by_name("search_VariantOnTranscript/Protein")
        searchbar.send_keys(search)
        searchbar.send_keys(Keys.RETURN)

        time.sleep(4)

        hit = False
        count = 0
        table = driver.find_element_by_xpath("//table[@id='viewlistTable_CustomVL_VOTunique_VOG_BRCA2']")
        for row in table.find_elements_by_xpath("//tr[@class='data']"):
            count = 0
            for table_data in row.find_elements_by_xpath(".//td"):
                count = count + 1
                if count == 6 and search in table_data.text:
                    # Searching on protein lev
                    row_cells = table_word.add_row().cells
                    row_cells[0].text = "LOVD"
                    row_cells[1].text = brca_input
                    row_cells[3].text = table_data.text
                    hit = True

                if count == 9 and hit == True:
                    row_cells[4].text = table_data.text
                    hit = False
                    print(table_data.text)

        p = document.add_paragraph("Results: ")    
        site_location = str(driver.current_url)
        add_hyperlink(p, 'Link', driver.current_url)

        driver.quit()

except Exception:
    pass

# Lindor pdf
# 4
# The script will read the pdf file containing the tables
# It will create a folder and paste the tables in excel files, this step is skipped if this folder already exists
# The script will iterate over each excel file in the folder

print("Searching Lindor pdf tables")
tables_pdf = read_pdf(folder_location + pdf_lindor, pages="17-28")
tables_output_location = "/tables excel lindor 2012"

tables_folder = folder_location + tables_output_location

if not os.path.exists(tables_folder):
    os.makedirs(tables_folder)

for i, table_pdf in enumerate(tables_pdf, start=1):
    table_pdf.to_excel(os.path.join(tables_folder, f"table_{i}.xlsx"), index=False)

# The excel files are divided in different groups based on the similarity of their structure (Equal amount of columns, columns that have the same meaning)
group_one = ('table_1.xlsx')
group_two = ("2.xlsx", "3.xlsx", "4.xlsx", "5.xlsx")
group_three = ("6.xlsx", "7.xlsx", "8.xlsx", "9.xlsx", "10.xlsx")
group_four = ("11.xlsx")
group_five = ("12.xlsx")

document.add_heading("Lindor 2012 pdf")

table = document.add_table(rows=1, cols=5)

table.style = 'Colorful List Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Gene'
hdr_cells[1].text = 'HGSV nucleotide'
hdr_cells[2].text = 'BIC nucleotide'
hdr_cells[3].text = "HGVS protein"
hdr_cells[4].text = 'Class'

tables_dir = os.listdir(tables_folder)

#In the first half it will search for BRCA1 genes

for table_file in tables_dir:
    if brca_input == "BRCA1":
        if table_file.endswith(group_one):

            wb = load_workbook(tables_folder + "/" + table_file)
            sheet = wb.active

            for row in sheet.iter_rows():
                for cell in row:
                    if search in str(cell.value):
                        row_cells = table.add_row().cells
                        row_cells[3].text = brca_input
                        row_cells[4].text = str(cell.offset(column=5).value)
                        #protein


        if table_file.endswith(group_two):

            wb = load_workbook(tables_folder + "/" + table_file)
            sheet = wb.active
            if sequence_type == "NUCLEOTIDE":

                for row in sheet.iter_rows(max_col = 3, min_col = 3):
                    for cell in row:
                        if search in str(cell.value):
                            row_cells = table.add_row().cells
                            row_cells[0].text = brca_input
                            row_cells[2].text = cell.value
                            row_cells[3].text = str(cell.offset(column=-1).value)
                            row_cells[4].text = str(cell.offset(column=5).value)
                            #bic
                        
                for row in sheet.iter_rows(max_col = 4, min_col = 4):
                    for cell in row:
                        if search in str(cell.value):
                            row_cells = table.add_row().cells
                            row_cells[0].text = brca_input
                            row_cells[1].text = cell.value
                            row_cells[2].text = str(cell.offset(column=-2).value)
                            row_cells[4].text = str(cell.offset(column=4).value)
                            #hgvs
            if sequence_type == "PROTEIN":
                for row in sheet.iter_rows(max_col = 2, min_col = 2):
                    for cell in row:
                        if search in str(cell.value):
                            row_cells = table.add_row().cells
                            row_cells[0].text = brca_input
                            row_cells[1].text = str(cell.offset(column=2).value)
                            row_cells[2].text = str(cell.offset(column=1).value)
                            row_cells[3].text = cell.value
                            row_cells[4].text = str(cell.offset(column=6).value)
                            #hgvs protein
                        
            
        if table_file.endswith(group_four):

            wb = load_workbook(tables_folder + "/" + table_file)
            sheet = wb.active
            for row in sheet.iter_rows():
                for cell in row:
                    if brca_input and search in str(cell.value):
                        row_cells = table.add_row().cells
                        row_cells[0].text = cell.value
                        row_cells[4].text = str(cell.offset(column=2).value)
                        #brca1/brca2

#Second half: it will search for the BRCA2 gene in the applicable file

    if brca_input == "BRCA2":
        if table_file.endswith(group_three):

            if sequence_type == "NUCLEOTIDE":
                wb = load_workbook(tables_folder + "/" + table_file)
                sheet = wb.active
                for row in sheet.iter_rows(max_col = 3, min_col = 3):
                    for cell in row:
                        if search in str(cell.value):
                            row_cells = table.add_row().cells
                            row_cells[0].text = brca_input
                            row_cells[2].text = cell.value
                            row_cells[3].text = str(cell.offset(column=-1).value)
                            row_cells[4].text = str(cell.offset(column=8).value)
                            #bic

                for row in sheet.iter_rows(max_col = 4, min_col = 4):
                    for cell in row:
                        if search in str(cell.value):
                            row_cells = table.add_row().cells
                            row_cells[0].text = brca_input
                            row_cells[1].text = cell.value
                            row_cells[3].text = str(cell.offset(column=-2).value)
                            row_cells[4].text = str(cell.offset(column=5).value)
                            #hgvs
            elif sequence_type == "PROTEIN":
                for row in sheet.iter_rows(max_col = 2, min_col = 2):
                    for cell in row:
                        if search in str(cell.value):
                            row_cells = table.add_row().cells
                            row_cells[0].text = brca_input
                            row_cells[1].text = str(cell.offset(column=2).value)
                            row_cells[2].text = str(cell.offset(column=1).value)
                            row_cells[3].text = cell.value
                            row_cells[4].text = str(cell.offset(column=9).value)
                            #hgvs protein


        if table_file.endswith(group_four):

            wb = load_workbook(tables_folder + "/" + table_file)
            sheet = wb.active
            for row in sheet.iter_rows():
                for cell in row:
                    if brca_input and search in str(cell.value):
                        row_cells = table.add_row().cells
                        row_cells[0].text = cell.value
                        row_cells[3].text = str(cell.offset(column=2).value)
                        
        if table_file.endswith(group_five):

            wb = load_workbook(tables_folder + "/" + table_file)
            sheet = wb.active
            for row in sheet.iter_rows(max_col = 4, min_col = 4):
                for cell in row:
                    if search in str(cell.value):
                        row_cells = table.add_row().cells
                        row_cells[0].text = brca_input
                        row_cells[1].text = cell.value
                        row_cells[3].text = str(cell.offset(column=5).value)
                        #hgvs

                     
            for row in sheet.iter_rows(max_col = 5, min_col = 5):
                for cell in row:
                    if search in str(cell.value):
                        row_cells = table.add_row().cells
                        row_cells[0].text = brca_input
                        row_cells[2].text = cell.value
                        row_cells[3].text = str(cell.offset(column=4).value)
                        #bic

#5) Reading excel sheet: bicbnl 27  fh 2-10-13_nov2013
# Python gives warning it can't parse the header but the script still works
# Similar to the other excel scripts except it differentiates between two tabs first.

print("Searching in the LOB excel file")
wb = load_workbook(folder_location + bicbnl)

document.add_heading("LOB")

table = document.add_table(rows=1, cols=5)
table.allow_outfit = True

table.style = 'Colorful List Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Gene'
hdr_cells[1].text = 'c.nom'
hdr_cells[2].text = 'Type'
hdr_cells[3].text = 'opm LOB'
hdr_cells[4].text = 'opmerking'
    
if brca_input == "BRCA1":
    # The script will open the right sheet first.
    sheet = wb['totaal BRCA1']

    for row in sheet.iter_rows(max_col = 5, min_col = 5):
        # It will search in the 5th column for the nucleotide
        for cell in row:
            if search in str(cell.value):
                row_cells = table.add_row().cells
                row_cells[0].text = brca_input
                row_cells[1].text = cell.value
                row_cells[2].text = str(cell.offset(column=2).value) # column G (type)
                row_cells[3].text = str(cell.offset(column=14).value) # column S (opmerking LOB indeling)
                row_cells[4].text = str(cell.offset(column=15).value) # column T (opmerking)

    for row in sheet.iter_rows(max_col = 6, min_col = 6):
        for cell in row:
            if search in str(cell.value):
                row_cells = table.add_row().cells
                row_cells[0].text = brca_input
                row_cells[1].text = cell.value
                row_cells[2].text = str(cell.offset(column=1).value)  # column G (type)
                row_cells[3].text = str(cell.offset(column=13).value)  # column S (opmerking LOB indeling)
                row_cells[4].text = str(cell.offset(column=14).value)  # column T (opmerking)

elif brca_input == "BRCA2":
    sheet = wb['totaal BRCA2']

    if sequence_type == "NUCLEOTIDE":
        for row in sheet.iter_rows(max_col = 5, min_col = 5):
            for cell in row:
                if search in str(cell.value):
                    row_cells = table.add_row().cells
                    row_cells[0].text = brca_input
                    row_cells[1].text = cell.value
                    row_cells[2].text = str(cell.offset(column=2).value)
                    row_cells[3].text = str(cell.offset(column=14).value)
                    row_cells[4].text = str(cell.offset(column=15).value)

    elif sequence_type == "PROTEIN":
        for row in sheet.iter_rows(max_col = 6, min_col = 6):
            for cell in row:
                if search in str(cell.value):
                    row_cells = table.add_row().cells
                    row_cells[0].text = brca_input
                    row_cells[1].text = cell.value
                    row_cells[2].text = str(cell.offset(column=1).value)  # column G (type)
                    row_cells[3].text = str(cell.offset(column=13).value)  # column S (opmerking LOB indeling)
                    row_cells[4].text = str(cell.offset(column=14).value)  # column T (opmerking)


#4/5/6 clinvar

print("Searching in the clinvar database")
eclient = Client(api_key="84c4d9fa268d5e41511ce024a0cf00537808")

gene_esearch = eclient.esearch(db='clinvar',term='{}[gene] AND {}'.format(brca_input, search))

id_list = gene_esearch.ids

document.add_heading("Clinvar")
table_word = document.add_table(rows=1, cols=4)
table_word.style = 'Colorful List Accent 1'
hdr_cells = table_word.rows[0].cells
hdr_cells[0].text = 'Gene'
hdr_cells[1].text = 'cDNA'
hdr_cells[2].text = 'Classification'
hdr_cells[3].text = 'Review Status'


for i in id_list:
    query = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi?db=clinvar&id={}&retmode=json".format(i)
    response = rq.get(query).json()
    id_string = str(i)
    if search in response["result"][id_string]["title"]:
        row_cells = table_word.add_row().cells
        row_cells[0].text = brca_input
        sequence = response["result"][id_string]["title"]
        sequence_split = sequence.split(":")
        row_cells[1].text = sequence_split[1]
        row_cells[2].text = response["result"][id_string]["clinical_significance"]["description"]
        if response["result"][id_string]["clinical_significance"]["review_status"] == "practice guideline":
            row_cells[3].text = "****"
        elif response["result"][id_string]["clinical_significance"]["review_status"] == "reviewed by expert panel":
            row_cells[3].text = "***"
        elif response["result"][id_string]["clinical_significance"]["review_status"] == "criteria provided, multiple submitters, no conflicts":
            row_cells[3].text = "**"
        elif response["result"][id_string]["clinical_significance"]["review_status"] == "criteria provided, conflicting interpretations" or 'criteria provided, single submitter':
            row_cells[3].text = "*"
        elif response["result"][id_string]["clinical_significance"]["review_status"] == "no assertion for the individual variant" or "no assertion criteria provided" or "no assertion provided":
            row_cells[3].text = "/"

current_url = "https://www.ncbi.nlm.nih.gov/clinvar/?term={}%5Bgene%5D+{}".format(brca_input, search)
p = document.add_paragraph("Results: ")    
site_location = current_url
add_hyperlink(p, 'Link', current_url)

# BIC
# Priority 6
print("Searching in the BIC text files")
folder_location = "C:/Users/jensv/Desktop/BRCA_-_prostate_cancer_and_ovarian_cancer_-_PARPi"
bic_brca1 = "/brca1_data_BIC.txt"
bic_brca2 = "/brca2_data_BIC.txt"

document.add_heading("BIC")

table = document.add_table(rows=1, cols=4)

table.style = 'Colorful List Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Gene'
hdr_cells[1].text = 'HGVS cDNA'
hdr_cells[2].text = "Protein"
hdr_cells[3].text = 'Category'


if brca_input == "BRCA1":
    full_path = folder_location + bic_brca1
    with open(full_path) as csv_file:
        csv_reader = csv.reader(csv_file, delimiter="\t")
        for row in csv_reader:
            count = 0
            hit = False
            for cell in row:
                count = count + 1
                if count == 8 and search in cell:
                    row_cells = table.add_row().cells
                    row_cells[0].text = brca_input
                    row_cells[1].text = cell
                    count = count + 1
                    hit = True

                elif count == 9 and search in cell:
                    row_cells = table.add_row().cells
                    row_cells[0].text = brca_input
                    row_cells[1].text = cell
                    count = count + 1
                    hit = True

                if count == 16 and hit == True:
                    row_cells[3].text = cell
                    hit = False

elif brca_input == "BRCA2":
    full_path = folder_location + bic_brca2
    with open(full_path) as csv_file:
        csv_reader = csv.reader(csv_file, delimiter="\t")
        for row in csv_reader:
            count = 0
            hit = False
            for cell in row:
                count = count + 1
                if count == 8 and search in cell:
                    row_cells = table.add_row().cells
                    row_cells[0].text = brca_input
                    row_cells[1].text = cell
                    count = count + 1
                    hit = True

                elif count == 9 and search in cell:
                    row_cells = table.add_row().cells
                    row_cells[0].text = brca_input
                    row_cells[1].text = cell
                    count = count + 1
                    hit = True

                if count == 16 and hit == True:
                    row_cells[3].text = cell
                    hit = False

#Enigma rules manual made table
print("Searching in the Enigma rules excel files")

enigma_table_brca1 = "/ENIGMA RULES BRCA1.xlsx"
enigma_table_brca2 = "/ENIGMA RULES BRCA2.xlsx"

if brca_input == "BRCA1":
    if sequence_type == "PROTEIN":
        wb = load_workbook(folder_location + enigma_table_brca1)
        sheet = wb.active
        document.add_heading("ENIGMA")
        table = document.add_table(rows=1, cols=6)

        table.style = 'Colorful List Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Gene'
        hdr_cells[1].text = 'Domain/motif'
        hdr_cells[2].text = 'AA start'
        hdr_cells[3].text = 'AA end'
        hdr_cells[4].text = 'AA alterations'
        hdr_cells[5].text = 'Classification'

        # searching for between AA start and AA end
        found = False
        for row in sheet.iter_rows(min_row=2, min_col=2, max_col=2):
            for cell in row:
                if int(search_mod) >= int(cell.value) and int(search_mod) <= int(cell.offset(column=1).value):
                    found = True
                    row_cells = table.add_row().cells
    
                    if search in cell.offset(column=2).value:
                        row_cells[0].text = brca_input
                        row_cells[1].text = str(cell.offset(column=-1).value)
                        row_cells[2].text = str(cell.value)
                        row_cells[3].text = str(cell.offset(column=1).value)
                        row_cells[4].text = search
                        row_cells[5].text = "Class 5"
                    else:
                        row_cells[0].text = brca_input
                        row_cells[1].text = str(cell.offset(column=-1).value)
                        row_cells[2].text = str(cell.value)
                        row_cells[3].text = str(cell.offset(column=1).value)
                        row_cells[4].text = search
                        row_cells[5].text = "Class 3"

    elif sequence_type == "NUCLEOTIDE":
        found = False
        for row in sheet.iter_rows(min_col=4, max_col=4):
            for cell in row:
                if search in str(cell.value) and found == False:
                    row_cells = table.add_row().cells
                    row_cells[0].text = brca_input
                    row_cells[1].text = str(cell.offset(column=-3).value)
                    row_cells[4].text = search
                    row_cells[5].text = "Class 5"

elif brca_input == "BRCA2":
    if sequence_type == "PROTEIN":
        wb = load_workbook(folder_location + enigma_table_brca2)
        sheet = wb.active
        document.add_heading("ENIGMA")
        table = document.add_table(rows=1, cols=6)

        table.style = 'Colorful List Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Gene'
        hdr_cells[1].text = 'Domain/motif'
        hdr_cells[2].text = 'AA start'
        hdr_cells[3].text = 'AA end'
        hdr_cells[4].text = 'AA alterations'
        hdr_cells[5].text = 'Classification'

        # searching for between AA start and AA end
        found = False
        for row in sheet.iter_rows(min_row=2, min_col=2, max_col=2):
            for cell in row:
                if int(search_mod) >= int(cell.value) and int(search_mod) <= int(cell.offset(column=1).value):
                    found = True
                    row_cells = table.add_row().cells

                    if search in cell.offset(column=2).value:
                        row_cells[0].text = brca_input
                        row_cells[1].text = str(cell.offset(column=-1).value)
                        row_cells[2].text = str(cell.value)
                        row_cells[3].text = str(cell.offset(column=1).value)
                        row_cells[4].text = search
                        row_cells[5].text = "Class 5"
                    else:
                        row_cells[0].text = brca_input
                        row_cells[1].text = str(cell.offset(column=-1).value)
                        row_cells[2].text = str(cell.value)
                        row_cells[3].text = str(cell.offset(column=1).value)
                        row_cells[4].text = search
                        row_cells[5].text = "Class 3"

    elif sequence_type == "NUCLEOTIDE":
        found = False
        for row in sheet.iter_rows(min_col=4, max_col=4):
            for cell in row:
                if search in str(cell.value) and found == False:
                    row_cells = table.add_row().cells
                    row_cells[0].text = brca_input
                    row_cells[1].text = str(cell.offset(column=-3).value)
                    row_cells[4].text = search
                    row_cells[5].text = "Class 5"

#2) Reading excel sheet : allenigmavariants_BICsubmission_2013-07-01
#7A

# The first column in this file is BRCA1/BRCA2, the second column is hgvsnucleotide and the fourth column is bic nucleotide
# The script will iterate over each row, if the brca_input matches with the gene in the 1st column AND the search is found in the 2nd OR 4th column it will give the class in column 7 as output

#cell.offset(column=2).value gives an error as output if the cell is empty but not when it is converted to string first.

print("Searching in the ENIGMA excel files")
wb = load_workbook(folder_location + variants_excel)

sheet = wb.active

document.add_heading("ENIGMA")

table = document.add_table(rows=1, cols=5)

table.style = 'Colorful List Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "List"
hdr_cells[1].text = 'Gene'
hdr_cells[2].text = 'HGSV nucleotide'
hdr_cells[3].text = 'HGSV protein'
hdr_cells[4].text = 'Class'


for row in sheet.iter_rows():
    for cell in row:
        # if BRCA gene is in 1st column AND nucleotide search is in 2ndcolumn (hgvs nucleotide)
        if brca_input in str(cell.value) and search in str(cell.offset(column=1).value) or brca_input in str(cell.value) and search in str(cell.offset(column=2).value):
            row_cells = table.add_row().cells
            row_cells[0].text = "Class 1 lijst met SNV's"
            row_cells[1].text = cell.value #gene
            row_cells[2].text = cell.offset(column=1).value #hgvs nucleotide
            row_cells[3].text = str(cell.offset(column=2).value) #hgvs protein
            row_cells[4].text = str(cell.offset(column=6).value) #class

#3)Reading excel sheet: Published Multifactoral data enigma
#7B samen met 7A
# This is likewise to searching in the allenigmavariants_BICsubmission_2013-07-01 excel file

print("Searching in the ENIGMA excel files part 2")
wb = load_workbook(folder_location + enigma)

sheet = wb.active

for row in sheet.iter_rows():
    for cell in row:
        if brca_input in str(cell.value) and search in str(cell.offset(column=1).value):
            row_cells = table.add_row().cells
            row_cells[0].text = "Enigma MLM lijst"
            row_cells[1].text = cell.value
            row_cells[2].text = cell.offset(column=1).value
            row_cells[3].text = str(cell.offset(column=2).value)
            row_cells[4].text = cell.offset(column=6).value
            Result = False
        if brca_input in str(cell.value) and search in str(cell.offset(column=2).value):
            row_cells = table.add_row().cells
            row_cells[0].text = "Enigma MLM lijst"
            row_cells[1].text = cell.value
            row_cells[2].text = cell.offset(column=1).value
            row_cells[3].text = str(cell.offset(column=2).value)
            row_cells[4].text = cell.offset(column=6).value

 #1) Reading excel sheet: 2020 functional categorization of BRCA1 VUS - CCR

print("Searching in the BRCA1 SNV lijst")
### Add heading to word document
document.add_heading("BRCA1 SNV lijst met functionele data (PMID: 32546644)", 1)

### Create table in word
table = document.add_table(rows=1, cols=4)

### Loading a style template and creating the header cells
table.style = 'Colorful List Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'DNA variant'
hdr_cells[1].text = 'Cisplatin Assay'
hdr_cells[2].text = 'Olaparib Assay'
hdr_cells[3].text = 'DR-GFP assay'


#If the input is BRCA1 the script will load the excel file. It will iterate over each row of the first column (max_col = 1) of the active sheet.
#If the search input is found in the first column the script will add the values of the 3rd, 4th and 5th column
#Example cell.offset(column=2).value means that the script will take the value of the 3rd column (1st column + 2 to the right = 3)
if brca_input == "BRCA1":

    wb = load_workbook(folder_location + functional_categorization)

    sheet = wb.active

    for row in sheet.iter_rows(max_col = 2, min_col = 2):
        for cell in row:
            if search in str(cell.value):
                    row_cells = table.add_row().cells
                    row_cells[0].text = cell.value
                    row_cells[1].text = cell.offset(column=1).value
                    row_cells[2].text = cell.offset(column=2).value
                    row_cells[3].text = cell.offset(column=3).value

 ########Franklin database
#9
print("Searching the Franklin database")
try:

    option = webdriver.ChromeOptions()
    option.add_argument('headless')
    PATH = "C:\Program Files (x86)\chromedriver.exe"
    driver = webdriver.Chrome(PATH, options = option)

    document.add_heading("Franklin/Varsome")
    table_word = document.add_table(rows=1, cols=3)
    table_word.style = 'Colorful List Accent 1'
    hdr_cells = table_word.rows[0].cells
    hdr_cells[0].text = "Engine"
    hdr_cells[1].text = 'Gene'
    hdr_cells[2].text = 'Classification'
    row_cells = table_word.add_row().cells

    driver.get("https://franklin.genoox.com/clinical-db/home")

    time.sleep(10)

    #a pop up appeared on this page which can be removed by pressing escape
    ActionChains(driver).send_keys(Keys.ESCAPE).perform()

    #The script will search for the searchbar and the button(somatic)
    searchbar = driver.find_element_by_xpath("/html/body/app-root/div/gnx-home-page/div/gnx-search/div      [2]/input")

    button = driver.find_element_by_xpath("/html/body/app-root/div/gnx-home-page/div/gnx-search/div[2]/     gnx-variant-type-toggle/button[2]")

    button.click()

    searchbar.send_keys(brca_input + ":" + search)
    searchbar.send_keys(Keys.RETURN)

    time.sleep(10)

    driver.find_element_by_xpath("//*[@id='mat-tab-label-2-3']").click()

    time.sleep(10)

    #classification = driver.find_element_by_xpath("//*[@id='mat-tab-content-2-3']/div/gnx-application-container/div/div/gnx-acmg-classification-app/gnx-result-page/div/gnx-result-top-section/div[2]/gnx-classifiction-gauge/div/div/svg/g/g/text[2]")
    classification = driver.find_element_by_xpath("//*[@class='indicator-text ng-star-inserted']")

    time.sleep(10)

    row_cells[0].text = "Franklin"
    row_cells[1].text = brca_input + ":" + search
    row_cells[2].text = classification.text
    franklin_class = classification.text
    print(classification.text)
    print(franklin_class)


    p = document.add_paragraph("Results: ")    
    site_location = str(driver.current_url)
    add_hyperlink(p, 'Link', driver.current_url)

    driver.quit()

except Exception:
    pass

#Varsome login required if too many searches are done
#10

try:

    option = webdriver.ChromeOptions()
    option.add_argument('headless')
    driver = webdriver.Chrome(PATH, options=option)

    driver.get("https://varsome.com/")

    time.sleep(1)

    searchbar = driver.find_element_by_xpath("//*[@id='search']")
    searchbar.send_keys(brca_input + " " + search)
    searchbar.send_keys(Keys.RETURN)

    time.sleep(1)

    #Popup asking to accept cookies 
    cookie = driver.find_element_by_xpath("//*[@id='onetrust-accept-btn-handler']")
    cookie.click()

    button = driver.find_element_by_xpath("//*[@id='proceedBtn']/h4")
    button.click()

    time.sleep(10)

    classification = driver.find_element_by_xpath("//*[@id='acmg_annotation']/div/div[1]/div/div[2]/div/    div/div[1]/div/div/span/div[2]/div/div")

    row_cells = table_word.add_row().cells
    row_cells[0].text = "Varsome"
    row_cells[1].text = brca_input + "" + search
    row_cells[2].text = classification.text

    p = document.add_paragraph("Results: ")    
    site_location = str(driver.current_url)
    add_hyperlink(p, 'Link', driver.current_url)
    print(driver.current_url)

    print("Searching Varsome database")

    driver.quit()

except Exception:
    pass

#Database of Functional Classifications of BRCA1 Variants based on Saturation Genome Editing
#

print("Searching database of functional classifications of BRCA1 Variants")
try:
    
    if brca_input == "BRCA1":

        option = webdriver.ChromeOptions()
        option.add_argument('headless')
        driver = webdriver.Chrome(PATH, options = option)

        document.add_heading("Database of Functional Classifications of BRCA1 Variants based on Saturation Genome Editing")
        table_word = document.add_table(rows=1, cols=3)
        table_word.style = 'Colorful List Accent 1'
        hdr_cells = table_word.rows[0].cells
        hdr_cells[0].text = 'Gene'
        hdr_cells[1].text = 'Transcript Variant'
        hdr_cells[2].text = 'Functional CLass'


        driver.get("https://sge.gs.washington.edu/BRCA1/#tab-9159-2")

        time.sleep(1)
        #The script will look up and click on the right tab
        tab = driver.find_element_by_xpath("//*[@id='main_page']/div/div/nav/div/ul/li[2]/a")
        tab.click()


        time.sleep(5)

        # look up the searchbar and enter the query
        searchbar = driver.find_element_by_xpath("//*[@id='DataTables_Table_0_filter']/label/input")
        searchbar.send_keys(search)
        searchbar.send_keys(Keys.RETURN)
    
        table = driver.find_element_by_xpath("//*[@id='DataTables_Table_0']/tbody")
 
        count = 0
        time.sleep(5)
        hit = False

        for row in table.find_elements_by_xpath("//tr"):
            count = 0
            for table_data in row.find_elements_by_xpath((".//td")):
                count = count + 1
                if count == 6 and search in table_data.text:
                    row_cells = table_word.add_row().cells
                    row_cells[0].text = brca_input
                    hit = True
                    row_cells[1].text = table_data.text 
                if count == 10 and hit == True:
                    row_cells[2].text = table_data.text
                    hit = False

    p = document.add_paragraph("Results: ")    
    site_location = str(driver.current_url)
    add_hyperlink(p, 'Link', driver.current_url)
    print(driver.current_url)

    driver.quit()
except Exception:
    pass

#Oncokb.org
# Using webscraping because the API is not free.

print("Searching Oncokb database")
try:

    option = webdriver.ChromeOptions()
    option.add_argument('headless')
    driver = webdriver.Chrome(PATH, options = option)

    driver.get("https://www.oncokb.org/")

    document.add_heading("OncoKB")
    table_word = document.add_table(rows=1, cols=2)
    table_word.style = 'Colorful List Accent 1'
    hdr_cells = table_word.rows[0].cells
    hdr_cells[0].text = 'Search'
    hdr_cells[1].text = 'Results'

    #time.sleep(5)

    searchbar = driver.find_element_by_xpath("/html/body/div[1]/div[2]/div[3]/div/div/div[3]/div/div/       div/div[1]/div[2]/div/input")
    searchbar.send_keys(brca_input + " " + search)
    time.sleep(2)
    searchbar.send_keys(Keys.RETURN)
    time.sleep(5)
    result = driver.find_element_by_xpath("//*[@id='root']/div[2]/div[3]/div/div/div/div[2]/div/div[2]")

    row_cells = table_word.add_row().cells
    row_cells[0].text = brca_input + search
    row_cells[1].text = result.text

    p = document.add_paragraph("Results: ")    
    site_location = str(driver.current_url)
    add_hyperlink(p, 'Link', driver.current_url)
    driver.quit()
    
except Exception:
    pass

# Pubmed
p = document.add_paragraph("pubmed: ") 
pubmed_site = "https://pubmed.ncbi.nlm.nih.gov/?term="  
pubmed_full = pubmed_site + brca_input + "+" + search
add_hyperlink(p, 'Link', pubmed_full)


document.save("output.docx")
print("Done")